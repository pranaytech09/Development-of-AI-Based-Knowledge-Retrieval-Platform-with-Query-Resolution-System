"""Extract and chunk text from uploaded PDF and Word documents."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
_CHUNK_SIZE = 1000
_CHUNK_OVERLAP = 100
_PDF_LIGATURES = str.maketrans(
    {
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
        "\ufb05": "st",
        "\ufb06": "st",
    }
)


def list_uploaded_files(uploads_dir: Path) -> list[Path]:
    """Return supported files in the uploads directory."""
    if not uploads_dir.is_dir():
        return []
    return sorted(
        path
        for path in uploads_dir.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def extract_text(path: Path) -> str:
    """Extract plain text from a PDF or Word file."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def chunk_text(text: str, chunk_size: int = _CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    normalized = _normalize_extracted_text(text)
    if not normalized:
        return []
    if len(normalized) <= chunk_size:
        return [normalized]

    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + chunk_size, len(normalized))
        chunks.append(normalized[start:end].strip())
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return [chunk for chunk in chunks if chunk]


def _normalize_extracted_text(text: str) -> str:
    """Normalize PDF ligatures and whitespace for consistent chunking and search."""
    return re.sub(r"\s+", " ", text.translate(_PDF_LIGATURES)).strip()


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError(f"Encrypted PDF is not supported: {path.name}")
    pages = [page.extract_text() or "" for page in reader.pages]
    return _normalize_extracted_text("\n".join(pages))


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [
        cell.text.strip()
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    return _normalize_extracted_text("\n".join([*paragraphs, *table_cells]))

